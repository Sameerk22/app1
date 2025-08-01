import os
import re
import time
import asyncio
import aiohttp
import streamlit as st
from tqdm.asyncio import tqdm_asyncio
from docx import Document
from tempfile import NamedTemporaryFile
from tqdm.asyncio import tqdm_asyncio
from pdf_utils import convert_pdf_to_docx_adobe  # <-- You must have this script locally

# --- CONFIGURATION ---
# OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
OPENROUTER_API_KEY = st.secrets["OPENROUTER_API_KEY"]
PDF_SERVICES_CLIENT_ID = st.secrets["PDF_SERVICES_CLIENT_ID"]
PDF_SERVICES_CLIENT_SECRET = st.secrets["PDF_SERVICES_CLIENT_SECRET"]

# MODEL = "gpt-4.1"
MAX_RETRIES = 3
RETRY_DELAY = 5
MAX_CONCURRENT_REQUESTS = 5

# --- Prompt Template ---
PROMPT_TEMPLATE = """
Here is a passage of a manuscript that we want to translate into {language} to make it easily readable for everyone. This is the original English version of one of the passages:
Can you rephrase this text in fluent and natural {language}? Please rewrite it and do not plagiarize. Do not give an answer, only give the translated text. 
Make sure that the reading experience “flows”, for example by not using the same words & sentence structures too often. Only translate, do not mention anything else. Please format the text properly without separating lines between the passages and exactly in the same structure it was so I can easily copy/paste it entirely into the manuscript of the book. You can remove numbers if that makes the reading experience better. You must translate/rewrite everything exactly and do not shorten it. Keep the quotes in their formatted way if applicable.
This is important for my career: Please try to translate everything sentence-by-sentence and do not make it shorter!
Important: Please Translate everything sentence-by-sentence in fluent {language} and do not make it shorter!
So do not skip translating any sentence line-by-line into {language} and do not make the translation shorter. This is crucial!
This is crucial: When you encounter general terminology such as "chapter," "the end,", “introduction” etc., or phrases that are widely used across books, do not vary or rephrase them—this is to ensure consistency. Always translate the text into the targetted language
Always use the most widely accepted or officially published book title, chapter name, or term if multiple synonyms are possible — e.g., use Meditations, not Reflections; Letters from a Stoic, not Stoic Letters. “Introduction”, not “getting started”
Also, always choose the most commonly published or recognized title, name or term in the target language — e.g., De Gedaanteverwisseling, not Metamorfose in Dutch. “Arme Peter” ipv “poor Peter” or “Armzalige Peter”
Important: Only give the exact translation; do not answer anything else, such as “Here’s a concise, modern rewrite,” etc.

Original Passage:
\"\"\"
{chunk}
\"\"\"
"""

# --- Helper Functions ---
def is_meaningful_text(text):
    cleaned = re.sub(r'[\W_]+', '', text)
    return bool(cleaned.strip())

def is_decorative_only(text):
    stripped = text.strip()
    return not stripped or re.fullmatch(r"[^\w\s]+", stripped) or re.fullmatch(r"[A-Z]", stripped)

# async def call_openai_gpt(session, prompt, semaphore):
#     url = "https://api.openai.com/v1/chat/completions"
#     headers = {
#         "Authorization": f"Bearer {OPENAI_API_KEY}",
#         "Content-Type": "application/json"
#     }
#     payload = {
#         "model": MODEL,
#         "messages": [{"role": "user", "content": prompt}],
#         "temperature": 0.7
#     }

#     for attempt in range(MAX_RETRIES):
#         async with semaphore:
#             try:
#                 async with session.post(url, headers=headers, json=payload) as resp:
#                     resp.raise_for_status()
#                     data = await resp.json()
#                     return data["choices"][0]["message"]["content"].strip()
#             except Exception as e:
#                 await asyncio.sleep(RETRY_DELAY)
#     return None


async def call_openai_gpt(session, prompt, semaphore):
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://yourdomain.com",  # Optional but recommended
        "X-Title": "EasyTranslate"                # Optional but helpful for tracking
    }
    payload = {
        "model": "anthropic/claude-sonnet-4",  # Update below based on your chosen model
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.1
    }

    for attempt in range(MAX_RETRIES):
        async with semaphore:
            try:
                async with session.post(url, headers=headers, json=payload) as resp:
                    resp.raise_for_status()
                    data = await resp.json()
                    return data["choices"][0]["message"]["content"].strip()
            except Exception as e:
                await asyncio.sleep(RETRY_DELAY)
    return None




async def call_with_progress(session, idx, prompt, semaphore, results, counter, lock, total, progress_callback):
    result = await call_openai_gpt(session, prompt, semaphore)
    results[idx] = result

    async with lock:
        counter[0] += 1
        pct = counter[0] / total
        if progress_callback:
            progress_callback(pct)

async def translate_docx_async(docx_path, output_path, language, progress_callback=None):
    doc = Document(docx_path)
    paragraphs = doc.paragraphs
    semaphore = asyncio.Semaphore(MAX_CONCURRENT_REQUESTS)

    jobs = []
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        original = para.text.strip()

        if re.fullmatch(r"[A-Z]", original) and i + 1 < len(paragraphs) and paragraphs[i + 1].text.strip()[:1].isupper():
            p = para._element
            p.getparent().remove(p)
            paragraphs = doc.paragraphs
            continue

        if not original or not is_meaningful_text(original) or is_decorative_only(original):
            i += 1
            continue


        word_count = len(original.split())
        is_heading = para.style.name.lower().startswith("heading") or para.alignment == 1

        # Skip short lines unless clearly uppercase or marked as a heading
        if word_count <= 1:
            if not original.isupper() and not is_heading:
                i += 1
                continue
            elif is_heading and language.lower() == "contemporary english":
                # Use a simple heading-specific prompt
                prompt = f"Rewrite this heading in contemporary English, keeping it short and clear:\n\"\"\"\n{original}\n\"\"\""
                jobs.append((i, para, prompt))
                i += 1
                continue


        prompt = PROMPT_TEMPLATE.format(chunk=original, language=language)
        jobs.append((i, para, prompt))
        i += 1

    total = len(jobs)
    results = [None] * total
    lock = asyncio.Lock()
    counter = [0]

    async with aiohttp.ClientSession() as session:
        await tqdm_asyncio.gather(
            *[
                call_with_progress(session, idx, prompt, semaphore, results, counter, lock, total, progress_callback)
                for idx, (_, _, prompt) in enumerate(jobs)
            ]
        )

    for (i, para, _), translated in zip(jobs, results):
        if translated:
            if translated.startswith('"""') and translated.endswith('"""'):
                translated = translated[3:-3].strip()
            elif translated.startswith('"""'):
                translated = translated[3:].strip()
            elif translated.endswith('"""'):
                translated = translated[:-3].strip()


            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = translated
            else:
                para.add_run(translated)

    doc.save(output_path)
    print(f"\n✅ Translated file saved to: {output_path}")


# --- Streamlit UI ---
st.title("📘 Easy Translate: Manuscript Translator")
uploaded_file = st.file_uploader("Upload a .docx or .pdf manuscript", type=["pdf", "docx"])
target_language = st.selectbox("Select output language", ["Contemporary English", "Spanish", "German", "Dutch"])

if uploaded_file:
    ext = os.path.splitext(uploaded_file.name)[-1].lower()

    with NamedTemporaryFile(delete=False, suffix=ext) as temp_input:
        temp_input.write(uploaded_file.read())
        temp_input.flush()
        input_path = temp_input.name

    # output_path = "translated_output.docx"
    base_name = os.path.splitext(uploaded_file.name)[0]  # Original filename without extension
    lang_suffix = target_language.lower().replace(" ", "_")  # e.g. "contemporary_english"
    final_filename = f"{base_name}_{lang_suffix}.docx"
    output_path = os.path.join(".", final_filename)  # Optional path control

    if ext == ".pdf":
        with st.spinner("🔄 Converting PDF to DOCX using Adobe..."):
            converted_docx = convert_pdf_to_docx_adobe(input_path)
            docx_path = converted_docx
    else:
        docx_path = input_path

    if st.button("🚀 Translate Now"):
        progress_bar = st.progress(0)

        def update_progress(pct):
            progress_bar.progress(pct)

        with st.spinner("Translating..."):
            asyncio.run(translate_docx_async(docx_path, output_path, target_language, update_progress))

        progress_bar.empty()
      
        st.success("Translation completed!")
        st.empty() 

        with open(output_path, "rb") as f:
            st.download_button("📥 Download Translated DOCX", f, file_name=final_filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
